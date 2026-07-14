"""
Агент анализа политических новостей Tengrinews.kz
===================================================

Что делает скрипт:
1. Заходит на раздел "Политика" сайта tengrinews.kz
2. Собирает статьи, вышедшие за последние N дней (по умолчанию 7)
3. Скачивает полный текст каждой статьи
4. Отправляет собранные тексты в Claude API для анализа
5. Сохраняет готовый отчёт в виде .docx файла с датой в названии
6. Запоминает дату последнего запуска, чтобы в следующий раз
   не анализировать одни и те же новости повторно

Как использовать:
    python tengrinews_analyzer.py

Для автозапуска раз в неделю без участия человека — см. инструкцию
в конце этого файла (или в сопроводительном README).
"""

import os
import re
import json
import time
from datetime import datetime, timedelta

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor


# ============================================================
# БЛОК 1. НАСТРОЙКИ
# ============================================================

# Базовый URL раздела "Политика" на Tengrinews
POLITICS_URL = "https://tengrinews.kz/tag/politic/"

# Сколько страниц пагинации листать за один прогон.
# Одна страница ~ 20 статей. При ежедневном запуске 2 страниц обычно достаточно.
PAGES_TO_SCAN = 2

# Если last_run.json ещё не создан (самый первый запуск) — за сколько
# последних дней брать новости. Дальше скрипт сам берёт "с прошлого запуска".
DAYS_WINDOW = 2

# Папка, куда будут сохраняться отчёты и служебные файлы
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
STATE_FILE = os.path.join(BASE_DIR, "last_run.json")

# Ключ API Google Gemini. Лучше не хранить его прямо в коде,
# а взять из переменной окружения (безопаснее).
# Как задать переменную окружения — см. инструкцию в конце файла.
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

# Модель, которой поручаем анализ (бесплатный тариф Gemini API)
GEMINI_MODEL = "gemini-2.5-flash"
GEMINI_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{GEMINI_MODEL}:generateContent"
)

# Заголовки для запросов, чтобы сайт не блокировал скрипт как бота
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}


# ============================================================
# БЛОК 2. СБОР СПИСКА СТАТЕЙ (СКРАПИНГ)
# ============================================================

def get_page_url(page_number: int) -> str:
    """Строит URL нужной страницы пагинации раздела Политика."""
    if page_number == 1:
        return POLITICS_URL
    return f"{POLITICS_URL}page/{page_number}/"


def parse_date(date_text: str) -> datetime | None:
    """
    Tengrinews показывает дату в формате '07 июля 16:17' или 'Сегодня 08:05'
    или 'Вчера 07:35'. Приводим к datetime.
    """
    date_text = date_text.strip().lower()
    now = datetime.now()

    if date_text.startswith("сегодня"):
        time_part = date_text.replace("сегодня", "").strip()
        h, m = map(int, time_part.split(":"))
        return now.replace(hour=h, minute=m, second=0, microsecond=0)

    if date_text.startswith("вчера"):
        time_part = date_text.replace("вчера", "").strip()
        h, m = map(int, time_part.split(":"))
        return (now - timedelta(days=1)).replace(hour=h, minute=m, second=0, microsecond=0)

    months = {
        "января": 1, "февраля": 2, "марта": 3, "апреля": 4,
        "мая": 5, "июня": 6, "июля": 7, "августа": 8,
        "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
    }
    match = re.match(r"(\d{1,2})\s+(\S+)\s+(\d{1,2}):(\d{2})", date_text)
    if not match:
        return None
    day, month_name, hour, minute = match.groups()
    month = months.get(month_name)
    if not month:
        return None
    year = now.year
    dt = datetime(year, month, int(day), int(hour), int(minute))
    # Если получилась дата в будущем (например, конец декабря/начало января),
    # значит статья на самом деле из прошлого года.
    if dt > now:
        dt = dt.replace(year=year - 1)
    return dt


def scan_politics_section(pages: int = PAGES_TO_SCAN) -> list[dict]:
    """
    Проходит по нескольким страницам раздела "Политика" и собирает
    заголовок, ссылку, короткое описание и дату каждой статьи.
    """
    articles = []

    for page in range(1, pages + 1):
        url = get_page_url(page)
        response = requests.get(url, headers=HEADERS, timeout=20)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # На странице статьи оформлены как блоки с заголовком h2/h3,
        # ссылкой на статью и датой рядом.
        for heading in soup.select("h2 a, h3 a"):
            href = heading.get("href")
            title = heading.get_text(strip=True)
            if not href or not title:
                continue
            if not href.startswith("http"):
                href = "https://tengrinews.kz" + href

            # Ищем дату — она обычно находится в соседнем элементе после заголовка
            date_node = heading.find_parent().find_next(string=re.compile(
                r"(Сегодня|Вчера|\d{1,2}\s+\S+\s+\d{1,2}:\d{2})", re.IGNORECASE
            ))
            article_date = parse_date(date_node) if date_node else None

            articles.append({
                "title": title,
                "url": href,
                "date": article_date,
            })

        time.sleep(1)  # вежливая пауза, чтобы не перегружать сайт запросами

    # убираем дубликаты по ссылке
    seen = set()
    unique_articles = []
    for art in articles:
        if art["url"] not in seen:
            seen.add(art["url"])
            unique_articles.append(art)

    return unique_articles


def filter_recent(articles: list[dict], since: datetime) -> list[dict]:
    """Оставляет только статьи, вышедшие после даты `since`."""
    return [a for a in articles if a["date"] and a["date"] >= since]


# ============================================================
# БЛОК 3. ЗАГРУЗКА ПОЛНОГО ТЕКСТА СТАТЬИ
# ============================================================

def fetch_article_text(url: str) -> str:
    """Скачивает страницу статьи и вытаскивает основной текст."""
    try:
        response = requests.get(url, headers=HEADERS, timeout=20)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Основной текст статьи на Tengrinews лежит в теге <article>
        # или в блоке с классом, содержащим "content"/"text"
        body = soup.find("article")
        if not body:
            body = soup.find("div", class_=re.compile("content|text|article"))
        if not body:
            return ""

        paragraphs = [p.get_text(" ", strip=True) for p in body.find_all("p")]
        text = "\n".join(p for p in paragraphs if p)
        return text[:4000]  # ограничиваем объём на случай очень длинных статей
    except Exception as e:
        print(f"  ! Не удалось загрузить текст статьи {url}: {e}")
        return ""


# ============================================================
# БЛОК 4. АНАЛИЗ ЧЕРЕЗ CLAUDE API
# ============================================================

def build_analysis_prompt(articles: list[dict]) -> str:
    """Собирает все статьи в один текстовый блок для передачи модели."""
    parts = []
    for i, art in enumerate(articles, 1):
        date_str = art["date"].strftime("%d.%m.%Y %H:%M") if art["date"] else "дата неизвестна"
        parts.append(
            f"[{i}] {art['title']} ({date_str})\n"
            f"Ссылка: {art['url']}\n"
            f"Текст: {art['text'] or '(текст недоступен)'}\n"
        )
    return "\n---\n".join(parts)


def analyze_with_gemini(articles: list[dict]) -> str:
    """Отправляет собранные статьи в Gemini и получает готовый аналитический отчёт."""
    if not GEMINI_API_KEY:
        raise RuntimeError(
            "Не найден ключ GEMINI_API_KEY. "
            "Задайте переменную окружения перед запуском скрипта."
        )

    articles_block = build_analysis_prompt(articles)

    system_prompt = (
        "Ты аналитик, который готовит ежедневные политические сводки "
        "на основе новостей Казахстана и мира с сайта Tengrinews.kz. "
        "Пиши по-русски, нейтрально и фактологично, без домыслов."
    )

    user_prompt = f"""
Ниже — список политических новостей за последний период с Tengrinews.kz.
Составь структурированный аналитический отчёт со следующими разделами:

1. Краткое резюме периода (3-5 предложений)
2. Ключевые темы и события (сгруппируй похожие новости)
3. Внутренняя политика Казахстана — главные события
4. Внешняя политика / международные новости, связанные с Казахстаном
5. Тенденции и на что стоит обратить внимание дальше

В конце добавь список использованных источников (заголовок + ссылка).

Новости:
{articles_block}
"""

    payload = {
        "system_instruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
        "generationConfig": {"maxOutputTokens": 4000},
    }

    response = requests.post(
        GEMINI_URL,
        headers={
            "Content-Type": "application/json",
            "x-goog-api-key": GEMINI_API_KEY,
        },
        json=payload,
        timeout=120,
    )
    response.raise_for_status()
    data = response.json()

    try:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError):
        raise RuntimeError(f"Неожиданный формат ответа Gemini: {data}")


# ============================================================
# БЛОК 5. СОХРАНЕНИЕ ОТЧЁТА В DOCX
# ============================================================

def save_report_docx(report_text: str, period_start: datetime, period_end: datetime) -> str:
    """Сохраняет отчёт как красиво оформленный Word-документ."""
    os.makedirs(REPORTS_DIR, exist_ok=True)

    doc = Document()

    title = doc.add_heading("Анализ политических новостей Tengrinews.kz", level=1)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        f"Период: {period_start.strftime('%d.%m.%Y')} — {period_end.strftime('%d.%m.%Y')}"
    )
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # пустая строка-отступ

    # Разбиваем текст отчёта по строкам и добавляем как параграфы.
    # Строки, начинающиеся с цифры и точки (например "1. ...") делаем заголовками.
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if re.match(r"^#{1,3}\s", line):
            doc.add_heading(re.sub(r"^#{1,3}\s", "", line), level=2)
        elif re.match(r"^\d+\.\s", line) and len(line) < 80:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    filename = f"tengrinews_politics_report_{period_end.strftime('%Y-%m-%d')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    return filepath


# ============================================================
# БЛОК 6. ХРАНЕНИЕ СОСТОЯНИЯ МЕЖДУ ЗАПУСКАМИ
# ============================================================

def load_last_run() -> datetime:
    """Читает дату предыдущего запуска. Если файла нет — берём 7 дней назад."""
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            return datetime.fromisoformat(data["last_run"])
    return datetime.now() - timedelta(days=DAYS_WINDOW)


def save_last_run(dt: datetime) -> None:
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump({"last_run": dt.isoformat()}, f, ensure_ascii=False, indent=2)


# ============================================================
# БЛОК 7. ГЛАВНАЯ ФУНКЦИЯ
# ============================================================

def main():
    run_started_at = datetime.now()
    since = load_last_run()

    print(f"Ищу политические новости с {since.strftime('%d.%m.%Y %H:%M')}...")
    all_articles = scan_politics_section(pages=PAGES_TO_SCAN)
    recent_articles = filter_recent(all_articles, since)

    if not recent_articles:
        print("Новых политических статей за этот период не найдено.")
        save_last_run(run_started_at)
        return

    print(f"Найдено статей: {len(recent_articles)}. Загружаю полные тексты...")
    for art in recent_articles:
        art["text"] = fetch_article_text(art["url"])
        time.sleep(1)

    print("Отправляю статьи на анализ в Gemini API...")
    report_text = analyze_with_gemini(recent_articles)

    print("Сохраняю отчёт в docx...")
    filepath = save_report_docx(report_text, since, run_started_at)

    save_last_run(run_started_at)
    print(f"Готово! Отчёт сохранён здесь: {filepath}")


if __name__ == "__main__":
    main()
